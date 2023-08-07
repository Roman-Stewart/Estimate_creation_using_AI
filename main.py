#!/usr/bin/env python
# coding: utf-8

# In[3]:


pip install pillow


# In[1]:


import tkinter as tk
from tkinter import filedialog
from PIL import Image, ImageTk
import shutil
import os

def open_image():
    file_path = filedialog.askopenfilename(filetypes=[("이미지 파일", "*.png;*.jpg;*.jpeg;*.gif")])
    if file_path:
        image = Image.open(file_path)
        image.thumbnail((300, 300))  # 이미지 크기 조절 (조정 가능)
        photo = ImageTk.PhotoImage(image)
        image_label.config(image=photo)
        image_label.image = photo
        message_label.config(text="이미지가 성공적으로 업로드되었습니다!")
        
        # 업로드한 이미지를 저장할 파일명 설정
        new_image_name = 'image1.jpg'
        
        # 데스크탑 경로 가져오기
        save_folder = r"C:/upload"
        os.makedirs(save_folder, exist_ok=True)
        
        # 파일을 저장
        destination_path = os.path.join(save_folder, new_image_name)
        shutil.copy(file_path, destination_path)
    else:
        message_label.config(text="이미지를 업로드하세요.")

# Tkinter 창 생성
root = tk.Tk()
root.title("이미지 업로드 및 저장")

# 이미지 업로드 버튼 생성
upload_button = tk.Button(root, text="이미지 업로드", command=open_image)
upload_button.pack(pady=10)

# 이미지를 표시할 라벨 생성
image_label = tk.Label(root)
image_label.pack(pady=10)

# 메시지를 표시할 라벨 생성
message_label = tk.Label(root, text="이미지를 업로드하세요.")
message_label.pack(pady=5)

root.mainloop()


# In[9]:


import os
from PIL import Image

def get_image_from_folder_on_desktop(folder_path, image_name):
    
    # 폴더가 존재하지 않으면 None 반환
    if not os.path.exists(folder_path):
        print(f"{folder_path} 경로가 데스크탑에 존재하지 않습니다.")
        return None
    
    # 이미지 파일 경로 생성
    image_path = os.path.join(folder_path, image_name)
    
    # 이미지 파일이 존재하지 않으면 None 반환
    if not os.path.exists(image_path):
        print(f"{image_name} 이미지 파일이 {folder_path} 경로에 존재하지 않습니다.")
        return None
    
    # 이미지 열기
    image = Image.open(image_path)
    return image

# 특정 폴더와 이미지 이름 지정
image_name = 'image1.jpg'
folder_path = r"C:/upload"
# 이미지 가져오기
image_variable = get_image_from_folder_on_desktop(folder_path, image_name)

# 이미지가 정상적으로 로드되면 사용 가능
if image_variable:
    # 이미지 처리 등 추가 작업 수행
    print("이미지를 성공적으로 가져왔습니다.")
    image_variable.show()  # 이미지를 보여줄 수 있습니다.
else:
    print("이미지를 가져오지 못했습니다.")


# In[10]:


image_variable


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[52]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# ## pytesseract 를 활용해 OCR하기

# In[5]:


pip install opencv-python pytesseract


# ## 파일과 답안지로 학습 데이터 만들기

# In[ ]:


# 예시로 사용할 이미지 파일과 텍스트 라벨을 리스트로 정의
image_paths = [
    'C:/training_images/image1.png',
    'C:/training_images/image2.png',
    'C:/training_images/image3.png',
    # 다른 이미지 파일의 경로를 추가로 정의할 수 있습니다.
]

text_labels = [
    '텍스트1',
    '텍스트2',
    '텍스트3',
    # 다른 텍스트 라벨을 추가로 정의할 수 있습니다.
]

# TSV 파일로 저장할 경로와 파일 이름 지정
tsv_file_path = 'C:/training_data/training_data.tsv'

# TSV 파일 생성 및 데이터 작성
with open(tsv_file_path, 'w', encoding='utf-8') as tsv_file:
    for image_path, text_label in zip(image_paths, text_labels):
        tsv_file.write(f'{image_path}\t{text_label}\n')

위 코드에서 사용하는 학습 데이터는 Tesseract-OCR Trainer에 로드하기 위해 TSV(Tab-Separated Values) 형식으로 구성되어야 합니다. TSV 파일은 각 줄마다 OCR 학습에 사용될 이미지 파일의 경로와 해당 이미지에 대한 텍스트 라벨이 탭으로 구분되어 있는 형식입니다.

TSV 파일의 구성은 다음과 같습니다:

python
Copy code
이미지파일1의경로    이미지1의텍스트라벨
이미지파일2의경로    이미지2의텍스트라벨
이미지파일3의경로    이미지3의텍스트라벨
...
이미지 파일의 경로와 해당 이미지에 대한 텍스트 라벨을 탭으로 구분하여 TSV 파일로 작성하면 됩니다. 예를 들어, 아래와 같이 training_data.tsv라는 파일로 학습 데이터를 구성할 수 있습니다:

javascript
Copy code
C:/training_images/image1.png    텍스트1
C:/training_images/image2.png    텍스트2
C:/training_images/image3.png    텍스트3
...
위 예시에서 이미지 파일의 경로와 텍스트 라벨은 탭으로 구분되어 있습니다. 이렇게 구성된 TSV 파일은 trainer.load_tsv() 메서드를 사용하여 학습에 사용할 수 있습니다.

OCR 모델을 학습시키기 위해 사용하는 학습 데이터는 이미지와 해당 이미지에 대한 텍스트 라벨의 쌍으로 구성되어야 하며, 이를 TSV 형식으로 작성하여 사용하시면 됩니다.
# ## 모델 학습

# 

# In[ ]:


import os
import cv2
import numpy as np
import tensorflow as tf
from tensorflow.keras import layers, models

# 데이터 경로 설정
train_data_dir = 'path_to_training_data'
test_data_dir = 'path_to_test_data'

# 데이터 불러오기 및 전처리 (필요한 경우)
train_data = ...
test_data = ...

# CNN 모델 설계
model = models.Sequential([
    layers.Conv2D(32, (3, 3), activation='relu', input_shape=(height, width, channels)),
    layers.MaxPooling2D((2, 2)),
    layers.Flatten(),
    layers.Dense(128, activation='relu'),
    layers.Dense(num_classes, activation='softmax')
])

# 모델 컴파일
model.compile(optimizer='adam',
              loss='sparse_categorical_crossentropy',
              metrics=['accuracy'])

# 모델 학습
model.fit(train_data, train_labels, epochs=2000, validation_data=(test_data, test_labels))

# 학습된 모델 저장
model.save('trained_ocr_model.h5')

# 추론 코드에서 모델 불러오기 및 사용
loaded_model = tf.keras.models.load_model('trained_ocr_model.h5')

# 이미지 전처리 및 추론
def preprocess_image(image_path):
    # 이미지 불러오기 및 전처리
    image = cv2.imread(image_path)
    # ... (이미지 전처리 작업)
    preprocessed_image = ...
    return preprocessed_image

def ocr_with_trained_model(image_path, model):
    preprocessed_image = preprocess_image(image_path)
    predicted_text = model.predict(preprocessed_image)
    return predicted_text

# OCR 수행
predicted_text = ocr_with_trained_model(image1_path, loaded_model)

# 출력
print("Predicted Text:", predicted_text)


# In[4]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[21]:





# In[20]:





# In[13]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# ## 학습한 모델을 이용해서 OCR

# In[1]:


import os
import cv2
import pytesseract

def ocr_korean_text_with_form(image_path, form_template_path):
    # Read form template image
    form_template = cv2.imread(form_template_path, 0)

    # Read image file into OpenCV
    image = cv2.imread(image_path)

    # Convert image to grayscale
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Apply adaptive thresholding for binarization
    _, binary_image = cv2.adaptiveThreshold(gray_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                         cv2.THRESH_BINARY, 11, 2)

    # Apply image median blur to remove noise
    blurred_image = cv2.medianBlur(binary_image, 3)

    # Text region detection (You may need to implement this step based on your specific form)
    # 예시로서, 이미지 전체 영역으로 텍스트 영역을 설정합니다.
    text_region = blurred_image.copy()

    # Extract text from text region
    custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1 -c preserve_interword_spaces=1'
    extracted_text = pytesseract.image_to_string(text_region, lang='kor', config=custom_config)

    return extracted_text

# specify a specific folder path
folder_path = r"C:/upload"

# image1 Specifies image file path
image1_filename = 'image1.jpg'
image1_path = os.path.join(folder_path, image1_filename)

# Form template image path
form_template_path = r"C:/form_template.jpg"

# Perform OCR on the image with the specific form
extracted_text = ocr_korean_text_with_form(image1_path, form_template_path)

# output extracted text
print(f"Text extracted from {image1_filename} with the form:")
print(extracted_text)


# In[6]:


print(extracted_text)


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# ## 견적 데이터 가져오기

# In[10]:


pip install pandas


# In[11]:


import pandas as pd

# 파일 경로 설정 (Iupload 폴더에 있는 data.xlsx 파일 경로)
data_route = r"C:/upload/data.xlsx"

# 데이터를 데이터 프레임으로 가져오기
df = pd.read_excel(data_route)

# 데이터 프레임 출력
df


# ## OCR 결과값 가정

# In[12]:


car_kind = "승용"
repair_list = ['컬러매칭 도장', '후드 교환', '콘솔박스 교환']


# ## OCR 결과값과 엑셀 견적내용 비교

# In[13]:


selected_car = df[df['차종'] == car_kind]
filtered_dataframe = selected_car[selected_car['견적내용'].isin(repair_list)]
filtered_dataframe


# ## 워드파일 채우기

# In[14]:


pip install python-docx


# In[20]:


## 차종 입력 (문자열 입력)
import os
from docx import Document

def fill_table_in_word(file_path, table_index, row_index, col_index, word):
    document = Document(file_path)
    
    # 특정 인덱스의 테이블 선택
    table = document.tables[table_index]
    
    # 특정 위치에 문자열 채우기
    table.cell(row_index, col_index).text = word

    # 변경된 내용 저장
    document.save(file_path)

# 워드 파일 경로
folder_path = r"C:/upload"
wordfile_name = "fill.docx"
file_path = os.path.join(folder_path, wordfile_name)

# 테이블 인덱스와 문자열 변수 word의 값을 채울 위치 설정
table_index = 0  # 첫 번째 테이블
row_index = 0    # 첫 번째 행
col_index = 4    # 다섯 번째 열
word = car_kind

# 기존의 워드 파일에서 특정 테이블에 데이터를 채우고 저장
fill_table_in_word(file_path, table_index, row_index, col_index, word)


# In[21]:


## 견적내용 입력 (데이터프레임)
from docx import Document
import pandas as pd

def fill_table_from_dataframe(file_path, table_index, start_row, start_col, dataframe):
    document = Document(file_path)
    
    # 특정 인덱스의 테이블 선택
    table = document.tables[table_index]

    # 데이터프레임의 데이터를 테이블에 채우기
    for row_index, row_data in enumerate(dataframe.values):
        for col_index, cell_value in enumerate(row_data):
            table.cell(row_index + start_row, col_index + start_col).text = str(cell_value)

    # 변경된 내용 저장
    document.save(file_path)

# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 2
col_index = 0

# 데이터프레임 예시
dataframe = filtered_dataframe.iloc[:, 1:2]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[22]:


## 부품비 입력

# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 2
col_index = 8

# 데이터프레임 예시
dataframe = filtered_dataframe.iloc[:, 2:3]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[23]:


## 공임비, 합계 입력

# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 2
col_index = 11

# 데이터프레임 예시
dataframe = filtered_dataframe.iloc[:, 3:5]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[24]:


## 합계 데이터프레임

sum_row = filtered_dataframe.select_dtypes(include='int').sum().to_frame().T
sum_row['부가가치세'] = int(sum_row['총합계비용'] * 0.1)
sum_row['부가가치세포함총액'] = sum_row['총합계비용'] + sum_row['부가가치세']
sum_row


# In[25]:


## 맨 밑 총정리 부품비 입력


# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 22
col_index = 1

# 데이터프레임 예시
dataframe = sum_row.iloc[:,0:1]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[26]:


## 맨 밑 총정리 공임비 입력


# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 22
col_index = 3

# 데이터프레임 예시
dataframe = dataframe = sum_row.iloc[:,1:2]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[27]:


## 맨 밑 총정리 계 입력


# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 22
col_index = 5

# 데이터프레임 예시
dataframe = dataframe = sum_row.iloc[:,2:3]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[28]:


## 맨 밑 총정리 부가가치세 입력

# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 22
col_index = 8

# 데이터프레임 예시
dataframe = dataframe = sum_row.iloc[:,3:4]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[29]:


## 맨 밑 총정리 총액 입력

# 테이블 인덱스 설정 (0부터 시작)
table_index = 1
row_index = 22
col_index = 10

# 데이터프레임 예시
dataframe = dataframe = sum_row.iloc[:,4:5]

# 데이터프레임의 데이터를 두 번째 테이블의 두 번째 행과 첫 번째 열부터 채우고 저장
fill_table_from_dataframe(file_path, table_index, start_row=row_index, start_col=col_index, dataframe=dataframe)


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:




